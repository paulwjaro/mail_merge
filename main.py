import docx

# Grabs the Guest List txt file
with open("Input/Invites/friend_invites.txt") as guest_list:
    new_guest_list = guest_list.read()

# Grabs the Sample Letter docx file
letter = docx.Document("./Input/Letters/starting_letter.docx")

# Splits the text doc of names into a list containing those names.
list_names = new_guest_list.split("\n")


# Takes the pre-made letter and converts the content into a String
new_text = []

for para in letter.paragraphs:
    new_text.append(para.text)

new_string = '\n'.join(new_text)

''''For every name in list names, replace the [name] text in the new string with the current name in the list. It then 
turns the string back into a list of paragraphs and creates a new docx file. It then adds each paragraph in the list to
the new document and saves the file as [name in letter]_letter.docx'''
for names in list_names:
    new_letter = new_string.replace("[name]", names)
    list_of_para = new_letter.split("\n")

    finished_letter = docx.Document()

    for para in list_of_para:
        finished_letter.add_paragraph(para)

    finished_letter.save(f"./Output/Ready To Send/{names}_letter.docx")
